from langchain_core.tools import tool
from docx import Document
from docx.shared import Inches, Pt
import os

import os
from docx import Document

STORAGE_PATH = os.getenv("DOCX_STORAGE_PATH")

def _path(docx_file_name: str) -> str:
    return os.path.join(STORAGE_PATH, docx_file_name)

@tool
def list_document_structure(docx_file_name: str):
    '''
    Tool to list all paragraphs in a docx file with their index, style, and text.
    '''
    doc = Document(_path(docx_file_name))
    return [
        {"index": i, "style": p.style.name, "text": p.text}
        for i, p in enumerate(doc.paragraphs)
    ]

@tool
def apply_title(title: str, docx_file_name: str):
    '''
    Tool to apply a title (Heading level 0) to a docx file.
    '''
    doc = Document(_path(docx_file_name))
    doc.add_heading(title, level=0)
    doc.save(_path(docx_file_name))


@tool
def edit_title(new_title: str, docx_file_name: str):
    '''
    Tool to edit the title in a docx file. Assumes the first paragraph is the title.
    '''
    doc = Document(_path(docx_file_name))
    if not doc.paragraphs:
        raise ValueError("No paragraphs found in the document.")
    doc.paragraphs[0].text = new_title
    doc.save(_path(docx_file_name))


@tool
def apply_header(header: str, docx_file_name: str, level: int = 1):
    '''
    Tool to apply a header (heading, level 1-9) to a docx file. Default level 1.
    '''
    doc = Document(_path(docx_file_name))
    doc.add_heading(header, level=level)
    doc.save(_path(docx_file_name))


@tool
def edit_header(paragraph_index: int, new_header: str, docx_file_name: str):
    '''
    Tool to edit a header (heading) at a given paragraph index in a docx file.
    '''
    doc = Document(_path(docx_file_name))
    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        raise IndexError("Paragraph index out of range.")
    if not paragraphs[paragraph_index].style.name.startswith("Heading"):
        raise ValueError(f"Paragraph {paragraph_index} is not a header.")
    paragraphs[paragraph_index].text = new_header
    doc.save(_path(docx_file_name))


@tool
def apply_paragraph(paragraph: str, docx_file_name: str):
    '''
    Tool to apply a paragraph to a docx file.
    '''
    doc = Document(_path(docx_file_name))
    doc.add_paragraph(paragraph)
    doc.save(_path(docx_file_name))


@tool
def edit_paragraph(paragraph_index: int, new_paragraph: str, docx_file_name: str):
    '''
    Tool to edit a paragraph in a docx file. Index is 0-based across ALL paragraphs
    (including title and headers), not just body text.
    '''
    doc = Document(_path(docx_file_name))
    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        raise IndexError("Paragraph index out of range.")
    paragraphs[paragraph_index].text = new_paragraph
    doc.save(_path(docx_file_name))

from docx.enum.text import WD_BREAK

@tool
def apply_page_break(docx_file_name: str):
    '''
    Tool to apply a page break to a docx file.
    '''
    doc = Document(_path(docx_file_name))
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    doc.save(_path(docx_file_name))

@tool
def apply_bullet_list(items: list[str], docx_file_name: str):
    '''
    Tool to apply a bulleted list to a docx file.
    '''
    doc = Document(_path(docx_file_name))
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    doc.save(_path(docx_file_name))


@tool
def apply_numbered_list(items: list[str], docx_file_name: str):
    '''
    Tool to apply a numbered list to a docx file.
    '''
    doc = Document(_path(docx_file_name))
    for item in items:
        doc.add_paragraph(item, style='List Number')
    doc.save(_path(docx_file_name))

@tool
def delete_paragraph(paragraph_index: int, docx_file_name: str):
    '''
    Tool to delete a paragraph at a given index from a docx file.
    '''
    doc = Document(_path(docx_file_name))
    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        raise IndexError("Paragraph index out of range.")
    p = paragraphs[paragraph_index]._element
    p.getparent().remove(p)
    doc.save(_path(docx_file_name))

TOOLS = [list_document_structure, apply_title, edit_title, apply_header, edit_header,apply_paragraph, edit_paragraph, apply_page_break, apply_bullet_list, apply_numbered_list, delete_paragraph]